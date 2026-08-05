"""Figures out of Word documents.

Same shape as the deck path — docling marks each picture's place, the extractor
yields pictures in that order, the two are paired positionally — with two
differences worth pinning.

Scope: the body only. Headers, footers and footnotes are separate parts that
docling does not mark, so a letterhead logo included here would shift every
pairing after it by one, silently.

Reach: anchored pictures count. `inline_shapes` sees only inline ones, and a
report figure with text wrapped around it is anchored — dropping those while
docling still marked their positions is the same off-by-one, arriving only for
documents laid out the way real reports are.
"""

from __future__ import annotations

import pytest

pytest.importorskip("docx", reason="python-docx arrives with the docling extra")

from docx import Document as DocxDocument  # noqa: E402
from docx.shared import Inches  # noqa: E402

from prismdoc.ooxml import extract_docx_images  # noqa: E402


def png(tmp_path, name: str, size: tuple[int, int] = (300, 200)):
    from PIL import Image

    path = tmp_path / name
    Image.new("RGB", size, (200, 40, 40)).save(path)
    return str(path)


class TestBodyPictures:
    def test_pictures_come_back_in_reading_order(self, tmp_path):
        first = png(tmp_path, "a.png", (300, 200))
        second = png(tmp_path, "b.png", (360, 200))

        document = DocxDocument()
        document.add_paragraph("Intro.")
        document.add_picture(first, width=Inches(3))
        document.add_paragraph("Middle.")
        document.add_picture(second, width=Inches(3))
        path = tmp_path / "doc.docx"
        document.save(path)

        images = extract_docx_images(str(path))

        assert [i.width for i in images] == [300, 360]

    def test_a_document_without_pictures_yields_none(self, tmp_path):
        document = DocxDocument()
        document.add_paragraph("Text only.")
        path = tmp_path / "plain.docx"
        document.save(path)

        assert extract_docx_images(str(path)) == []

    def test_page_index_is_zero_because_docx_has_no_pages(self, tmp_path):
        """Better than inventing a number that means nothing downstream."""
        document = DocxDocument()
        document.add_picture(png(tmp_path, "a.png"), width=Inches(3))
        path = tmp_path / "doc.docx"
        document.save(path)

        image = extract_docx_images(str(path))[0]

        assert image.page_index == 0
        assert image.bbox is None

    def test_the_bytes_are_the_stored_asset(self, tmp_path):
        """No rasterising — what comes out is what the document stored."""
        document = DocxDocument()
        document.add_picture(png(tmp_path, "a.png", (321, 123)), width=Inches(1))
        path = tmp_path / "doc.docx"
        document.save(path)

        image = extract_docx_images(str(path))[0]

        assert (image.width, image.height) == (321, 123)
        assert image.mime == "image/png"
        assert image.data.startswith(b"\x89PNG")


class TestScope:
    def test_a_header_logo_is_not_a_figure(self, tmp_path):
        """Docling marks no position for it, so counting it shifts every pairing."""
        document = DocxDocument()
        document.add_paragraph("Body text.")
        document.add_picture(png(tmp_path, "body.png"), width=Inches(3))
        document.sections[0].header.paragraphs[0].add_run().add_picture(
            png(tmp_path, "logo.png"), width=Inches(1)
        )
        path = tmp_path / "letterhead.docx"
        document.save(path)

        assert len(extract_docx_images(str(path))) == 1

    def test_an_anchored_picture_is_found(self, tmp_path):
        """The case `inline_shapes` misses — text wrapped around a figure."""
        document = DocxDocument()
        document.add_paragraph("Intro.")
        document.add_picture(png(tmp_path, "a.png"), width=Inches(3))
        path = tmp_path / "anchored.docx"
        document.save(path)

        # Convert the inline drawing into an anchored one, as a wrapped figure is.
        reopened = DocxDocument(str(path))
        inline = reopened.element.body.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )
        assert inline, "fixture produced no drawing to anchor"

        images = extract_docx_images(str(path))
        assert len(images) == 1, "a wrapped figure would be dropped by inline_shapes"
